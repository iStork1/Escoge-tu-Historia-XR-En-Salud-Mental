import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def is_table_separator(line: str) -> bool:
    """Check if a line is a markdown table separator."""
    stripped = line.strip()
    if not stripped.startswith('|') or not stripped.endswith('|'):
        return False
    cells = [cell.strip() for cell in stripped.split('|')[1:-1]]
    return all(all(c in '-: ' for c in cell) for cell in cells)

def is_table_row(line: str) -> bool:
    """Check if a line is a markdown table row."""
    return line.strip().startswith('|') and line.strip().endswith('|')

def parse_table_row(line: str) -> list:
    """Parse a markdown table row into cells."""
    stripped = line.strip()
    cells = [cell.strip() for cell in stripped.split('|')[1:-1]]
    return cells

def add_table_from_lines(doc, lines: list, start_idx: int) -> int:
    """
    Add a table to the document from markdown lines.
    Returns the index of the line after the table.
    """
    table_lines = []
    idx = start_idx
    
    # Collect all table rows (header + separator + data rows)
    while idx < len(lines) and is_table_row(lines[idx]):
        if not is_table_separator(lines[idx]):
            table_lines.append(lines[idx])
        idx += 1
    
    if not table_lines:
        return start_idx
    
    # Parse the table
    rows = [parse_table_row(line) for line in table_lines]
    if not rows:
        return idx
    
    num_cols = len(rows[0])
    
    # Create table in docx
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    
    # Fill in the table
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            # Clear default paragraph
            cell.text = ''
            # Add content
            p = cell.paragraphs[0]
            p.text = cell_text
            # Format header row (first row)
            if row_idx == 0:
                for run in p.runs:
                    run.font.bold = True
    
    return idx

def md_to_docx(md_path: Path, docx_path: Path):
    text = md_path.read_text(encoding='utf-8')
    lines = text.splitlines()
    doc = Document()
    
    idx = 0
    while idx < len(lines):
        line = lines[idx]
        
        # Check for table
        if is_table_row(line):
            idx = add_table_from_lines(doc, lines, idx)
        elif line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
            idx += 1
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
            idx += 1
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
            idx += 1
        elif line.startswith('+ '):
            doc.add_paragraph(line[2:].strip(), style='List Bullet')
            idx += 1
        elif line.startswith('- '):
            # Check if this is part of a table (avoid treating | as bullet)
            if '|' in line:
                idx += 1
                continue
            doc.add_paragraph(line[2:].strip(), style='List Bullet')
            idx += 1
        elif line.startswith('```'):
            # naive code block handling: add as plain paragraph
            doc.add_paragraph(line)
            idx += 1
        elif line.strip():  # Non-empty line
            doc.add_paragraph(line)
            idx += 1
        else:  # Empty line
            idx += 1
    
    doc.save(str(docx_path))

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print('Usage: python scripts/generate_docx_from_md.py <input.md> <output.docx>')
        sys.exit(1)
    md = Path(sys.argv[1])
    out = Path(sys.argv[2])
    if not md.exists():
        print('Input MD not found:', md)
        sys.exit(2)
    md_to_docx(md, out)
    print('Saved DOCX to:', out)
