from docx import Document
import os
md_path = r"c:\Users\Felipe Jaimes\Desktop\Escoje Tu Historia XR En Salud Mental\instructions\8-prompt-decisions.md"
docx_path = r"c:\Users\Felipe Jaimes\Desktop\Escoje Tu Historia XR En Salud Mental\instructions\8-prompt-decisions.docx"

doc = Document()

# Read the markdown and add lines as paragraphs; keep code blocks intact
with open(md_path, 'r', encoding='utf-8') as f:
    lines = f.read().splitlines()

for line in lines:
    # Preserve empty lines
    if line.strip() == '':
        doc.add_paragraph('')
    else:
        doc.add_paragraph(line)

# Ensure output directory exists
out_dir = os.path.dirname(docx_path)
if not os.path.isdir(out_dir):
    os.makedirs(out_dir)

# Save the document
doc.save(docx_path)
print(f"Saved DOCX to: {docx_path}")
